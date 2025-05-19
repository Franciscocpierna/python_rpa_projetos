from docx import Document
from docx.shared import Pt

#Abrindo o arquivo do Word
#arquivoWord = Document("Certificado1.docx")
arquivoWord = Document(r"c:\python_projetos\python_rpa_projetos\python_word\Certificado1.docx")

#Configurando os estilos
estilo = arquivoWord.styles["Normal"]

#for = para
for paragrafo in arquivoWord.paragraphs:
    #if = se
    if "@nome" in paragrafo.text:
        paragrafo.text = "Amanda Batista Alves"
        fonte = estilo.font
        fonte.name = "Calibri (Corpo)"
        fonte.size = Pt(24)
        

#Salva o certificado com o nome do aluno
arquivoWord.save("Amanda Batista Alves.docx")

print("Certificado gerado com sucesso!")