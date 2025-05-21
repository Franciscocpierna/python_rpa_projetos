from docx import Document

#Aumenta o tamanho da fonte
from docx.shared import Pt

#Mudar a cor da fonte
from docx.shared import RGBColor

#Abre o arquivo de excel
from openpyxl import load_workbook


#Configurando o caminho
nome_arquivo_alunos = "C:\\python_projetos\\python_rpa_projetos\\python_word\\DadosAlunos.xlsx"

#Abrindo o arquivo
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)

#Selecionando a aba sheet/planilha
sheet_selecionada = planilhaDadosAlunos["Nomes"]

#for = para
for linha in range(2, len(sheet_selecionada["A"]) + 1):

    #Abrindo o arquivo do Word
    arquivoWord = Document("C:\\python_projetos\\python_rpa_projetos\\python_word\\Certificado2.docx")

    #Configurando os estilos
    estilo = arquivoWord.styles["Normal"]
    
    #Pegamos o nome do aluno que está na linha corrente
    nomeAluno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    nomeCurso = sheet_selecionada['E%s' % linha].value
    nomeInstrutor = sheet_selecionada['F%s' % linha].value
    
    frase_parte1 = "Concluiu com sucesso o curso de "
    frase_parte2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
    frase_montada = f"{frase_parte2} {dia} de {mes} de {ano}"

    #for = para
    for paragrafo in arquivoWord.paragraphs:
        #if = se
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            
        if "Dezembro" in paragrafo.text:
            paragrafo.text = frase_parte1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso) #Adiciona o nome do curso no final do paragrafo
            adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0) #Muda a cor para vermelho
            adicionaNovaPalavra.underline = True #Subrinha a palavra
            adicionaNovaPalavra.bold = True #Deixa a palavra como negrito
            adicionaNovaPalavra = paragrafo.add_run(frase_montada) #Adiciona o texto final do paragrafo
            adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0) #Muda a cor para preto
            
        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhoCertificado = "C:\\python_projetos\\python_rpa_projetos\\python_word\\" + nomeAluno + ".docx"

    #Salva o certificado com o nome do aluno
    arquivoWord.save(caminhoCertificado)

print("Certificados gerados com sucesso!")